# app/services/export_service.py - ROBUST EXPORT SERVICE WITH GRACEFUL DEPENDENCIES
"""
Export service that handles missing dependencies gracefully and doesn't break blueprint imports
"""
import io
import os
import json
import tempfile
import zipfile
from datetime import datetime
from typing import Dict, List, Optional, Union, BinaryIO
from flask import current_app
import logging

logger = logging.getLogger(__name__)

# Graceful dependency imports - don't break if libraries are missing
REPORTLAB_AVAILABLE = False
PYTHON_DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
    logger.info("ReportLab available - PDF export enabled")
except ImportError:
    logger.info("ReportLab not available - PDF export disabled")

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    PYTHON_DOCX_AVAILABLE = True
    logger.info("python-docx available - DOCX export enabled")
except ImportError:
    logger.info("python-docx not available - DOCX export disabled")

class ExportService:
    """Service for exporting stories to various formats with graceful dependency handling"""
    
    def __init__(self):
        """Initialize export service with available formats"""
        self.supported_formats = ['txt', 'html', 'json']
        
        if REPORTLAB_AVAILABLE:
            self.supported_formats.append('pdf')
        
        if PYTHON_DOCX_AVAILABLE:
            self.supported_formats.append('docx')
        
        logger.info(f"ExportService initialized with formats: {self.supported_formats}")
    
    def get_supported_formats(self) -> List[str]:
        """Get list of currently supported export formats"""
        return self.supported_formats.copy()
    
    def is_format_supported(self, format: str) -> bool:
        """Check if a format is supported"""
        return format.lower() in self.supported_formats
    
    def export_story(self, project, scenes: List, format: str = 'txt') -> BinaryIO:
        """
        Export a complete story to the specified format
        
        Args:
            project: Project model instance
            scenes: List of scene model instances
            format: Export format ('txt', 'html', 'pdf', 'docx', 'json')
        
        Returns:
            BinaryIO: File-like object containing the exported story
        
        Raises:
            ValueError: If format is not supported
            RuntimeError: If export fails
        """
        format = format.lower()
        
        if not self.is_format_supported(format):
            available = ', '.join(self.supported_formats)
            raise ValueError(f"Format '{format}' not supported. Available formats: {available}")
        
        try:
            if format == 'txt':
                return self._export_txt(project, scenes)
            elif format == 'html':
                return self._export_html(project, scenes)
            elif format == 'json':
                return self._export_json(project, scenes)
            elif format == 'pdf' and REPORTLAB_AVAILABLE:
                return self._export_pdf(project, scenes)
            elif format == 'docx' and PYTHON_DOCX_AVAILABLE:
                return self._export_docx(project, scenes)
            else:
                # Fallback to txt if specific format fails
                logger.warning(f"Format {format} failed, falling back to TXT")
                return self._export_txt(project, scenes)
                
        except Exception as e:
            logger.error(f"Export failed for format {format}: {str(e)}")
            raise RuntimeError(f"Export failed: {str(e)}")
    
    def _export_txt(self, project, scenes: List) -> BinaryIO:
        """Export story as plain text"""
        output = io.StringIO()
        
        # Header
        output.write(f"{project.title}\n")
        output.write("=" * len(project.title) + "\n\n")
        
        if project.description:
            output.write(f"{project.description}\n\n")
        
        # Metadata
        output.write(f"Genre: {project.genre or 'Unspecified'}\n")
        output.write(f"Target Audience: {project.target_audience or 'General'}\n")
        output.write(f"Current Word Count: {project.current_word_count or 0:,}\n")
        if project.target_word_count:
            output.write(f"Target Word Count: {project.target_word_count:,}\n")
        output.write(f"Exported: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}\n\n")
        
        output.write("-" * 50 + "\n\n")
        
        # Scenes
        for i, scene in enumerate(sorted(scenes, key=lambda s: s.order_index or 0), 1):
            output.write(f"Scene {i}: {scene.title}\n")
            output.write("-" * (len(f"Scene {i}: {scene.title}")) + "\n\n")
            
            if scene.description:
                output.write(f"Description: {scene.description}\n\n")
            
            if scene.content:
                output.write(f"{scene.content}\n\n")
            
            output.write("\n" + "=" * 30 + "\n\n")
        
        # Convert to bytes
        text_content = output.getvalue()
        output.close()
        
        return io.BytesIO(text_content.encode('utf-8'))
    
    def _export_html(self, project, scenes: List) -> BinaryIO:
        """Export story as HTML"""
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{project.title}</title>
    <style>
        body {{
            font-family: 'Georgia', 'Times New Roman', serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        .header {{
            text-align: center;
            border-bottom: 2px solid #ccc;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        .title {{
            font-size: 2.5em;
            margin-bottom: 10px;
            color: #2c3e50;
        }}
        .metadata {{
            background: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 30px;
        }}
        .scene {{
            margin-bottom: 40px;
            border-left: 4px solid #3498db;
            padding-left: 20px;
        }}
        .scene-title {{
            font-size: 1.8em;
            color: #2c3e50;
            margin-bottom: 10px;
        }}
        .scene-description {{
            font-style: italic;
            color: #666;
            margin-bottom: 15px;
        }}
        .scene-content {{
            text-align: justify;
            white-space: pre-wrap;
        }}
        .export-info {{
            font-size: 0.9em;
            color: #666;
            text-align: center;
            margin-top: 40px;
            border-top: 1px solid #ccc;
            padding-top: 20px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1 class="title">{project.title}</h1>
        {f'<p>{project.description}</p>' if project.description else ''}
    </div>
    
    <div class="metadata">
        <strong>Project Information:</strong><br>
        Genre: {project.genre or 'Unspecified'}<br>
        Target Audience: {project.target_audience or 'General'}<br>
        Current Word Count: {project.current_word_count or 0:,}<br>
        {f'Target Word Count: {project.target_word_count:,}<br>' if project.target_word_count else ''}
        Status: {project.status or 'Active'}<br>
        Phase: {project.current_phase or 'Development'}
    </div>
"""
        
        # Add scenes
        for i, scene in enumerate(sorted(scenes, key=lambda s: s.order_index or 0), 1):
            html_content += f"""
    <div class="scene">
        <h2 class="scene-title">Scene {i}: {scene.title}</h2>
        {f'<div class="scene-description">{scene.description}</div>' if scene.description else ''}
        {f'<div class="scene-content">{scene.content or "No content yet."}</div>' if scene.content else ''}
    </div>
"""
        
        html_content += f"""
    <div class="export-info">
        Exported from ALVIN on {datetime.utcnow().strftime('%Y-%m-%d at %H:%M UTC')}
    </div>
</body>
</html>"""
        
        return io.BytesIO(html_content.encode('utf-8'))
    
    def _export_json(self, project, scenes: List) -> BinaryIO:
        """Export story as JSON"""
        export_data = {
            'export_metadata': {
                'version': '1.0',
                'exported_at': datetime.utcnow().isoformat(),
                'exported_by': 'ALVIN v1.0',
                'format': 'json'
            },
            'project': {
                'title': project.title,
                'description': project.description,
                'genre': project.genre,
                'target_audience': project.target_audience,
                'expected_length': project.expected_length,
                'status': project.status,
                'current_phase': project.current_phase,
                'current_word_count': project.current_word_count,
                'target_word_count': project.target_word_count,
                'tone': project.tone,
                'estimated_scope': project.estimated_scope,
                'marketability': project.marketability,
                'original_idea': project.original_idea,
                'created_at': project.created_at.isoformat() if project.created_at else None,
                'updated_at': project.updated_at.isoformat() if project.updated_at else None
            },
            'scenes': [
                {
                    'title': scene.title,
                    'description': scene.description,
                    'content': scene.content,
                    'scene_type': scene.scene_type,
                    'emotional_intensity': scene.emotional_intensity,
                    'order_index': scene.order_index,
                    'status': scene.status,
                    'location': scene.location,
                    'conflict': scene.conflict,
                    'hook': scene.hook,
                    'character_focus': scene.character_focus,
                    'word_count': scene.word_count,
                    'dialog_count': scene.dialog_count,
                    'created_at': scene.created_at.isoformat() if scene.created_at else None,
                    'updated_at': scene.updated_at.isoformat() if scene.updated_at else None
                }
                for scene in sorted(scenes, key=lambda s: s.order_index or 0)
            ],
            'statistics': {
                'total_scenes': len(scenes),
                'total_word_count': sum(scene.word_count or 0 for scene in scenes),
                'average_scene_length': sum(scene.word_count or 0 for scene in scenes) // max(len(scenes), 1),
                'scene_types': list(set(scene.scene_type for scene in scenes if scene.scene_type))
            }
        }
        
        json_content = json.dumps(export_data, indent=2, ensure_ascii=False)
        return io.BytesIO(json_content.encode('utf-8'))
    
    def _export_pdf(self, project, scenes: List) -> BinaryIO:
        """Export story as PDF (requires reportlab)"""
        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("PDF export requires reportlab library")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        story.append(Paragraph(project.title, title_style))
        
        # Description
        if project.description:
            story.append(Paragraph(project.description, styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Metadata
        metadata_text = f"""
        <b>Genre:</b> {project.genre or 'Unspecified'}<br/>
        <b>Target Audience:</b> {project.target_audience or 'General'}<br/>
        <b>Current Word Count:</b> {project.current_word_count or 0:,}<br/>
        {f'<b>Target Word Count:</b> {project.target_word_count:,}<br/>' if project.target_word_count else ''}
        <b>Exported:</b> {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}
        """
        story.append(Paragraph(metadata_text, styles['Normal']))
        story.append(PageBreak())
        
        # Scenes
        for i, scene in enumerate(sorted(scenes, key=lambda s: s.order_index or 0), 1):
            # Scene title
            scene_title_style = ParagraphStyle(
                'SceneTitle',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12
            )
            story.append(Paragraph(f"Scene {i}: {scene.title}", scene_title_style))
            
            # Scene description
            if scene.description:
                story.append(Paragraph(f"<i>{scene.description}</i>", styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Scene content
            if scene.content:
                # Split content into paragraphs
                paragraphs = scene.content.split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        story.append(Paragraph(paragraph.strip(), styles['Normal']))
                        story.append(Spacer(1, 12))
            
            story.append(Spacer(1, 24))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def _export_docx(self, project, scenes: List) -> BinaryIO:
        """Export story as DOCX (requires python-docx)"""
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError("DOCX export requires python-docx library")
        
        doc = Document()
        
        # Title
        title = doc.add_heading(project.title, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Description
        if project.description:
            doc.add_paragraph(project.description)
            doc.add_paragraph()
        
        # Metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run('Genre: ').bold = True
        metadata_para.add_run(project.genre or 'Unspecified')
        metadata_para.add_run('\nTarget Audience: ').bold = True
        metadata_para.add_run(project.target_audience or 'General')
        metadata_para.add_run('\nCurrent Word Count: ').bold = True
        metadata_para.add_run(f"{project.current_word_count or 0:,}")
        if project.target_word_count:
            metadata_para.add_run('\nTarget Word Count: ').bold = True
            metadata_para.add_run(f"{project.target_word_count:,}")
        metadata_para.add_run('\nExported: ').bold = True
        metadata_para.add_run(datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC'))
        
        doc.add_page_break()
        
        # Scenes
        for i, scene in enumerate(sorted(scenes, key=lambda s: s.order_index or 0), 1):
            # Scene title
            scene_heading = doc.add_heading(f"Scene {i}: {scene.title}", level=2)
            
            # Scene description
            if scene.description:
                desc_para = doc.add_paragraph()
                desc_para.add_run(scene.description).italic = True
                doc.add_paragraph()
            
            # Scene content
            if scene.content:
                # Split content into paragraphs
                paragraphs = scene.content.split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            
            doc.add_paragraph()
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

# Create singleton instance
export_service = ExportService()